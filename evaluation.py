import os
from docx import Document
import RbzRate
import json
import RbzRate
import Operation
import VatOption
Item = 1
companyName = 'Takudzwa'

try:
	filename = 'RbzRate.json'
	Vat_value = 0.15
	productCount =0
	quantity =0
	with open(filename) as read_file:
		Rbz_rate = json.load(read_file)
	print('\nWelcome To Telone Evaluation Calculator\nPress any letter to Exit Any Time ')
	RbzRate.Rate()
	products = int(input('Number Of Products To be Evaluated: '))
	if products == 1:
		quantity = int(input('Enter Quantity: '))
		while True:
			VatSelection = VatOption.VatOptions()
			if VatSelection == '3':
				print('Thanks for Using Our Application')
				os.sys.exit()
			Operation.operation(VatSelection,quantity,Rbz_rate,products,productCount)
			# elif VatSelection == '1' or VatSelection =='2':
			# 	UnitPrice = float(input('Enter Supplier Unit Price: '))
			# 	if VatSelection == '1':
			# 		print('VAT INCLUDED')
			# 		TotalPrice =   UnitPrice * quantity
			# 		UnitPriceUSD = UnitPrice / Rbz_rate
			# 		TotalPriceUSD = UnitPriceUSD * quantity
			# 		print('Unit Price(ZWL)		 Unit Price(USD)		TotalPrice(zwl)')
			# 		print( f'{str(UnitPrice)}0 		  {str(UnitPriceUSD)}0 	{str(TotalPrice)}0' )
			# 	elif VatSelection == '2':
			# 		UnitPriceVat =UnitPrice+(UnitPrice * Vat_value) 
			# 		TotalPrice = UnitPriceVat * quantity
			# 		UnitPriceUSD = UnitPriceVat / Rbz_rate
			# 		TotalPriceUSD = TotalPrice / Rbz_rate
			# 		# prices = [['Item','companyName',UnitPrizeVat,UnitPriceUSD,TotalPrice]]
			# 		print('VAT EXCLUDED')
			# 		print('Unit Price(ZWL)		Unit Price(USD)			TotalPrice(zwl)')
			# 		print( f'{str(UnitPriceVat)}0 		{str(UnitPriceUSD)}0 	{str(TotalPrice)}0' )
			# else:
			# 	print(' Invalid Option Selected') 	

	elif products >= 2:
		while True:
			if productCount >= 1:
				productCount=0
				products = int(input('Number Of Products To be Evaluated: '))
			elif productCount == 0:
				VatSelection = VatOption.VatOptions()
				if VatSelection == '3':
					print('Thanks for Using Our Application')
					os.sys.exit()
				while  productCount < products :
					document = Document()
					#Operation.operation(VatSelection,quantity,Rbz_rate,products,productCount)
					if VatSelection == '1' or VatSelection == '2':
						quantity = int(input('Enter Quantity: '))
						UnitPrice = float(input('Supplier Unit Price: '))
						# document.add_heading('5.0 Finincial Evaluation')
						# table = document.add_table(rows=6, cols=6)
						# table_header = [ 'Item','Company','Quantity','UnitPrizeZWL','UnitUSDPrice','TotalPrice']
						if VatSelection == '1':
							TotalPrice =   (UnitPrice * quantity)
							UnitPriceUSD = UnitPrice / Rbz_rate
							TotalPriceUSD = TotalPrice / Rbz_rate
							#prices = [[Item,companyName,UnitPrizeZWL,UnitPriceUSD,TotalPrice]]
							# for i in range(6):
							# 	hdr_cells = table.rows[0].cells[i].text = table_header[i]
							# for  Items,companyName,Quantity,UnitPrizeZWL,UnitUSDPrice,TotalPrice in prices:
							# 		row_cells = table.add_row().cells
							# 		row_cells[0].text = str(Item)
							# 		row_cells[2].text =  str(Quantity)
							# 		row_cells[3].text =str(UnitPrizeZWL)
							# 		row_cells[4].text = str(UnitUSDPrice)
							# 		row_cells[5].text =str(TotalPrice)
							# 		row_cells[1].text = companyName
						
							print('\nVAT INCLUDED')
							print('Unit Price(ZWL)		 Unit Price(USD)		TotalPrice(ZWL)			TotalPrice(USD)')
							print( f'{str(UnitPrice)}0 		  {str(UnitPriceUSD)}0 	{str(TotalPrice)}0			{str(TotalPriceUSD)}0\n' )
							
							productCount = productCount + 1
						elif VatSelection == '2':
							UnitPriceVat =UnitPrice+(UnitPrice * Vat_value) 
							TotalPrice = UnitPriceVat * quantity
							UnitPriceUSD = UnitPriceVat / Rbz_rate
							TotalPriceUSD = TotalPrice / Rbz_rate
							# prices = [['Item','companyName',UnitPrizeVat,UnitPriceUSD,TotalPrice]]
							print('\nVAT INCLUDED')
							print('Unit Price(ZWL)		 Unit Price(USD)		TotalPrice(ZWL)			TotalPrice(USD)')
							print( f'{str(UnitPriceVat)}0 		  {str(UnitPriceUSD)}0 	{str(TotalPrice)}0			{str(TotalPriceUSD)}0\n' )
							productCount = productCount + 1
					
					else:
						print(' Invalid Option Selected') 
						break
			else:
				pass
	else:
		print('Sorry Products should Not Be less Than 1!')
				
		# document.save('Financial.docx')		
    
except ValueError:
    print('Thanks for Using This Application')
except FileNotFoundError:
    print('File Not Found') 
		
		
			
	