import os
from docx import Document
import docx
Vat_value = 0.15
productCount = 0
Item = 1
companyName = 'Takudzwa'
try:
	def operation(VatOption,quantity,Rbz_rate,products,productCount):
		if VatOption == '1' or VatOption =='2':
			# if products > 1:
			# 	quantity = int(input('Enter Quantity: '))	
			UnitPrice = float(input('Enter Supplier Unit Price: '))
			if VatOption == '1':
				TotalPrice =   UnitPrice * quantity
				UnitPriceUSD = UnitPrice / Rbz_rate
				TotalPriceUSD = UnitPriceUSD * quantity
				print('\nVAT INCLUDED')
				print('Unit Price(ZWL)		 Unit Price(USD)		TotalPrice(ZWL)			TotalPrice(USD)')
				print( f'{str(UnitPrice)}0 		  {str(UnitPriceUSD)}0 	{str(TotalPrice)}0			{str(TotalPriceUSD)}0\n' )
				if products > 1:
					productCount = productCount + 1
			elif VatOption == '2':
				UnitPriceVat =UnitPrice+(UnitPrice * Vat_value) 
				TotalPrice = UnitPriceVat * quantity
				UnitPriceUSD = UnitPriceVat / Rbz_rate
				TotalPriceUSD = TotalPrice / Rbz_rate
				# prices = [['Item','companyName',UnitPrizeVat,UnitPriceUSD,TotalPrice]]
				print('\nVAT EXCLUDED')
				print('Unit Price(ZWL)		 Unit Price(USD)		TotalPrice(ZWL)			TotalPrice(USD)')
				print( f'{str(UnitPriceVat)}0 		  {str(UnitPriceUSD)}0 	{str(TotalPrice)}0			{str(TotalPriceUSD)}0\n' )
				# if products > 1:
				# 	productCount = productCount + 1
			document = docx.Document()
			document.add_heading('5.0 Finincial Evaluation')	 
			table_header = [ 'Item','Company','Quantity','UnitPrizeZWL','UnitUSDPrice','TotalPrice']
			records =[
			[Item,companyName,quantity,UnitPrice,UnitPriceUSD,TotalPrice],
			[Item,companyName,quantity,UnitPrice,UnitPriceUSD,TotalPrice]]
			
			table = document.add_table(rows=6, cols=6)
			table.style = 'Table Grid'
			hdr_cells = table.rows[0].cells
			hdr_cells[0].text = 'Item'
			hdr_cells[1].text = 'Company'
			hdr_cells[2].text =  'Quantity'
			hdr_cells[3].text ='UnitPrize'
			hdr_cells[4].text = 'UnitUSDPrice'
			hdr_cells[5].text ='TotalPrice'
			for Items, UnitUSDPrice in records:
				
				row_cells = table.add_row().cells
				# row_cells[0].text = str(Item)
				# #row_cells[1].text = companyName
				# row_cells[2].text =  str(quantity)
				# row_cells[3].text =str(UnitPrize)
				row_cells[0].text = str(UnitPriceUSD)
				# row_cells[5].text =str(TotalPrice)
			document.save('Finincial.docx')
			os.system('start  Finincial.docx')

			print('now4')
			
				
		else:
			print(' Invalid Option Selectedd')

		return operation
except:
	pass
